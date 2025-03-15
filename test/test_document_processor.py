"""
Test script for the DocumentProcessor module
"""

import os
from utils.document_processor import DocumentProcessor
from docx import Document

def create_test_files():
    """Create sample test files for demonstration"""
    # Create a sample directory for test files
    os.makedirs("data/sample", exist_ok=True)
    
    # Create a sample Word document for testing
    doc_path = "data/sample/sample.docx"
    doc = Document()
    doc.add_heading('Sample Document', 0)
    doc.add_paragraph('This is a sample document for testing the document processor.')
    doc.add_paragraph('It contains multiple paragraphs to demonstrate text extraction.')
    doc.add_heading('Section 1', level=1)
    doc.add_paragraph('This is the content of section 1. The document processor should extract this text and split it into chunks.')
    doc.add_heading('Section 2', level=1)
    doc.add_paragraph('This is the content of section 2. LangChain document loaders are used to process this document.')
    doc.save(doc_path)
    
    print(f"Created sample Word document at {doc_path}")
    print("Note: For actual PDF and XLSX files, please place them in the data directory")

def test_document_processor():
    """Test the DocumentProcessor class functionality"""
    processor = DocumentProcessor(chunk_size=500, chunk_overlap=100)
    
    print("DocumentProcessor initialized with chunk_size=500, chunk_overlap=100")
    print("\nTo process files, use:")
    print("1. processor.process_file('path/to/document.pdf')")
    print("2. processor.process_file('path/to/document.docx')")
    print("3. processor.process_file('path/to/spreadsheet.xlsx')")
    print("\nTo process all documents in a directory:")
    print("processor.process_directory('path/to/directory')")
    
    # Check if any supported files exist in the data directory
    data_dir = "data"
    if os.path.exists(data_dir):
        supported_files = []
        for root, _, files in os.walk(data_dir):
            for file in files:
                if file.endswith(('.pdf', '.docx', '.xlsx', '.xls')):
                    supported_files.append(os.path.join(root, file))
        
        if supported_files:
            print("\nFound supported files:")
            for file in supported_files:
                print(f"- {file}")
            
            # Try to process a DOCX file if available
            docx_files = [f for f in supported_files if f.endswith('.docx')]
            if docx_files:
                test_file = docx_files[0]
                print(f"\nTesting document processing with {test_file}...")
                try:
                    chunks = processor.process_file(test_file)
                    print(f"Successfully processed {test_file}")
                    print(f"Extracted {len(chunks)} chunks")
                    if chunks:
                        print("\nSample chunk content:")
                        print(chunks[0][:200] + "..." if len(chunks[0]) > 200 else chunks[0])
                except Exception as e:
                    print(f"Error processing file: {str(e)}")
            
            print("\nYou can process these files using the DocumentProcessor")
        else:
            print("\nNo supported files found in the data directory")
            print("Please add PDF, DOCX, or XLSX files to test with actual documents")

if __name__ == "__main__":
    create_test_files()
    test_document_processor()
